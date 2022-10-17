import pandas as pd 
from docx import Document 
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#  STEP 1: READ DATA FROM EXCEL FILE
PDE_list = pd.ExcelFile("Filename.xlsx")
PDE_data_frame = PDE_list.parse("Form Responses 1") 
aList = PDE_data_frame["Column1"].tolist()
bList = PDE_data_frame["Column2"].tolist()
newlist = zip(aList, bList) 
# print(set(newlist))

# STEP 2: DEFINE FUNCTION
def create_word_doc(name, university): 

    font.size = Pt(60) 
    par1 = document.add_paragraph(name) 
    par1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    font.size = Pt(32)
    par2 = document.add_paragraph(university)
    par2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()
    return par1, par2

# STEP 3: ITERATE OVER PARTICIPANTS LIST AND PRINT TO WORD
document = Document()
Tstyle = document.styles['Normal']
font = Tstyle.font
for participant_name, participant_university in newlist:
    head, subhead = create_word_doc(participant_name, participant_university)

document.save('Participant List Iowa PDE Seminar.docx')


